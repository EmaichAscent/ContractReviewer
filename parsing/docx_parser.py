"""Parse .docx files using python-docx."""

from docx import Document


def parse_docx(filepath):
    """Extract structured text from a .docx file.

    Returns a list of dicts with keys: style, text, index
    """
    doc = Document(filepath)
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            paragraphs.append({
                "style": para.style.name if para.style else "Normal",
                "text": text,
                "index": i,
            })

    # Also extract table content
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append({
                    "style": "Table",
                    "text": row_text,
                    "index": f"table-{ti}-row-{ri}",
                })

    return paragraphs


def get_full_text(filepath):
    """Get the full text of a .docx as a single string."""
    paragraphs = parse_docx(filepath)
    return "\n".join(p["text"] for p in paragraphs)


def detect_file_type(filepath):
    """Detect file type by magic bytes. Returns 'docx', 'doc', 'pdf', or 'unknown'."""
    with open(filepath, "rb") as f:
        header = f.read(5)
    if header[:4] == b"PK\x03\x04":
        return "docx"
    elif header[:4] == b"\xd0\xcf\x11\xe0":
        return "doc"
    elif header[:5] == b"%PDF-":
        return "pdf"
    return "unknown"
