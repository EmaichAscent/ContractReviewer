"""Generate a standalone Word document of proposed contract revisions.

This is the artifact that lets the user share a clean change-list with their
client or legal counsel. Unlike the marked-up client contract (which requires
the original to be a .docx), this document is generated from the analysis
JSON and works for any input format — PDFs included.

For each criterion that scored below "Strong", the document shows:
  - The current contract language (if any was quoted in the analysis)
  - The proposed replacement / insertion language
  - The criterion's explanation and recommended location

It's organized by scoring category so the reader can see which weaknesses
cluster together (e.g., everything under "Risk Transference").
"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_revisions_document(analysis, client_name, output_path, jurisdiction=None):
    """Build the Suggested Contract Revisions .docx.

    Args:
        analysis: the analysis dict produced by analyze_contract
        client_name: client display name for the header
        output_path: target .docx path
        jurisdiction: optional dict with state/state_abbrev
    """
    doc = Document()

    _add_header(doc, client_name, analysis, jurisdiction)
    _add_intro(doc)
    weak_count = _add_category_sections(doc, analysis)
    _add_gap_analysis(doc, analysis)
    _add_statute_section(doc, analysis)
    _add_footer(doc, weak_count)

    doc.save(output_path)


def _add_header(doc, client_name, analysis, jurisdiction):
    title = doc.add_heading("Suggested Contract Revisions", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run("Client: ").bold = True
    meta.add_run(f"{client_name}\n")
    if jurisdiction and jurisdiction.get("state"):
        abbrev = jurisdiction.get("state_abbrev", "")
        meta.add_run("Jurisdiction: ").bold = True
        meta.add_run(f"{jurisdiction['state']}" + (f" ({abbrev})" if abbrev else "") + "\n")
    overall = analysis.get("overall_score", 0)
    meta.add_run("Overall Score: ").bold = True
    meta.add_run(f"{int(overall * 100)}%\n")
    meta.add_run("Review Date: ").bold = True
    meta.add_run(datetime.now().strftime("%Y-%m-%d"))


def _add_intro(doc):
    intro = doc.add_paragraph()
    intro.add_run(
        "This document summarizes the recommended revisions to the management "
        "agreement based on a clause-by-clause review from the management "
        "company's perspective. Each suggested change includes the current "
        "contract language (when available) and proposed replacement or "
        "insertion text. Use this as a negotiation worksheet or hand-off to "
        "legal counsel."
    )


def _add_category_sections(doc, analysis):
    """Add a section per scoring category, listing each weak criterion's
    proposed revision. Returns the total count of revisions written."""
    categories = analysis.get("categories", {})
    total_written = 0

    for cat_name, cat_data in categories.items():
        if not isinstance(cat_data, dict):
            continue

        # Collect criteria that need revision (score < 2)
        weak = []
        for crit_id, crit in cat_data.get("criteria", {}).items():
            if isinstance(crit, dict) and crit.get("score", 2) < 2:
                weak.append((crit_id, crit))

        if not weak:
            continue

        cat_score = cat_data.get("score", 0)
        doc.add_heading(f"{cat_name} ({int(cat_score * 100)}%)", level=1)

        if cat_data.get("summary"):
            doc.add_paragraph(cat_data["summary"])

        for crit_id, crit in weak:
            _add_criterion_revision(doc, crit_id, crit)
            total_written += 1

    return total_written


def _add_criterion_revision(doc, crit_id, crit):
    """Render one criterion's proposed revision."""
    crit_name = _humanize_criterion_id(crit_id)
    doc.add_heading(crit_name, level=2)

    score = crit.get("score", 0)
    level_label = "MISSING" if score == 0 else "WEAK"
    status = doc.add_paragraph()
    status_run = status.add_run(f"Current Status: {level_label}")
    status_run.bold = True
    status_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) if score == 0 else RGBColor(0xA8, 0x55, 0x00)

    if crit.get("explanation"):
        doc.add_paragraph(crit["explanation"])

    if crit.get("contract_text"):
        label = doc.add_paragraph()
        label.add_run("Current Contract Language:").bold = True
        quote = doc.add_paragraph()
        quote.paragraph_format.left_indent = Inches(0.4)
        quote_run = quote.add_run(f'"{crit["contract_text"]}"')
        quote_run.italic = True
    else:
        absent = doc.add_paragraph()
        absent_run = absent.add_run("[No corresponding language in current contract — this protection is missing entirely.]")
        absent_run.italic = True

    if crit.get("suggested_revision"):
        label = doc.add_paragraph()
        label.add_run("Proposed Language:").bold = True
        proposed = doc.add_paragraph()
        proposed.paragraph_format.left_indent = Inches(0.4)
        proposed.add_run(crit["suggested_revision"])

    if crit.get("revision_location"):
        loc = doc.add_paragraph()
        loc_run = loc.add_run(f"Recommended Location: {crit['revision_location']}")
        loc_run.italic = True
        loc_run.font.size = Pt(10)


def _add_gap_analysis(doc, analysis):
    """Render the gap-analysis section (favorable client language worth importing)."""
    gaps = analysis.get("gap_analysis", [])
    if not gaps:
        return

    doc.add_page_break()
    doc.add_heading("Gap Analysis — Client Language Worth Importing into the Master Template", level=1)
    doc.add_paragraph(
        "The following language appears in the client contract and strengthens "
        "the management company's position, but is not present in the firm's "
        "master template. Consider importing these provisions for future "
        "engagements."
    )

    for gap in gaps:
        if not isinstance(gap, dict):
            continue
        section_label = gap.get("client_section") or "Provision"
        doc.add_heading(section_label, level=2)
        if gap.get("client_quote"):
            quote = doc.add_paragraph()
            quote.paragraph_format.left_indent = Inches(0.4)
            quote.add_run(f'"{gap["client_quote"]}"').italic = True
        if gap.get("template_section"):
            p = doc.add_paragraph()
            p.add_run("Suggested template section: ").bold = True
            p.add_run(gap["template_section"])
        if gap.get("rationale"):
            p = doc.add_paragraph()
            p.add_run("Why import: ").bold = True
            p.add_run(gap["rationale"])


def _add_statute_section(doc, analysis):
    """Render the statutory compliance notes section."""
    concerns = analysis.get("statute_concerns", [])
    if not concerns:
        return

    doc.add_page_break()
    doc.add_heading("Statutory Compliance Notes", level=1)
    doc.add_paragraph(
        "Specific conflicts between the contract language and applicable state "
        "statutes. Each item identifies the contract section, the statutory "
        "requirement, and the legal/operational impact."
    )
    for concern in concerns:
        if isinstance(concern, str):
            doc.add_paragraph(concern, style="List Bullet")


def _add_footer(doc, weak_count):
    doc.add_page_break()
    summary = doc.add_paragraph()
    summary.add_run(f"Total revisions proposed: ").bold = True
    summary.add_run(f"{weak_count}")

    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run(
        "This document is a contract analysis worksheet produced by Contract "
        "Reviewer. It is intended to support negotiation and legal review and "
        "does not constitute legal advice. Final contract terms should be "
        "reviewed by qualified counsel licensed in the governing jurisdiction."
    )
    disclaimer_run.italic = True
    disclaimer_run.font.size = Pt(9)
    disclaimer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _humanize_criterion_id(crit_id):
    """Strip category prefixes (profit_, empower_, etc.) and snake_case → Title Case."""
    cleaned = crit_id
    for prefix in ("profit_", "empower_", "risk_", "board_", "value_", "statute_", "compliance_"):
        if cleaned.startswith(prefix):
            cleaned = cleaned[len(prefix):]
            break
    return cleaned.replace("_", " ").title()
