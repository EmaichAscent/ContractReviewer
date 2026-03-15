"""Generate polished client-ready scorecard as .docx."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime


def generate_scorecard(analysis_results, client_name, output_path, jurisdiction=None):
    """Generate a professional scorecard document.

    Args:
        analysis_results: dict from claude_analyzer.analyze_contract()
        client_name: name of the client association
        output_path: where to save the .docx
        jurisdiction: jurisdiction info dict
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    title = doc.add_heading("CONTRACT REVIEW SCORECARD", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)

    # Client info
    doc.add_paragraph()
    info_para = doc.add_paragraph()
    info_para.add_run("Client Name: ").bold = True
    info_para.add_run(client_name)
    info_para.add_run("\nReview Date: ").bold = True
    info_para.add_run(datetime.now().strftime("%B %d, %Y"))
    if jurisdiction and jurisdiction.get("state"):
        info_para.add_run("\nJurisdiction: ").bold = True
        info_para.add_run(f"{jurisdiction['state']} ({jurisdiction.get('state_abbrev', '')})")

    doc.add_paragraph()

    # Overall Score
    overall_score = analysis_results.get("overall_score", 0)
    overall_pct = round(overall_score * 100)
    score_heading = doc.add_heading(f"Overall Score: {overall_pct}%", level=1)
    _color_heading(score_heading, _score_color(overall_score))

    doc.add_paragraph()

    # Score summary table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Evaluation Area"
    hdr_cells[1].text = "Score"
    hdr_cells[2].text = "Rating"

    # Bold the header row
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    categories = analysis_results.get("categories", {})
    for cat_name, cat_data in categories.items():
        score = cat_data.get("score", 0)
        row = table.add_row()
        row.cells[0].text = cat_name
        row.cells[1].text = f"{round(score * 100)}%"
        row.cells[2].text = _score_rating(score)

    doc.add_paragraph()

    # Detailed evaluation for each category
    doc.add_heading("Evaluation Details", level=1)
    doc.add_paragraph()

    for cat_name, cat_data in categories.items():
        score = cat_data.get("score", 0)
        heading = doc.add_heading(f"{cat_name} — {round(score * 100)}%", level=2)
        _color_heading(heading, _score_color(score))

        # Category summary
        summary = cat_data.get("summary", "")
        if summary:
            doc.add_paragraph(summary)

        # Individual criteria results
        criteria = cat_data.get("criteria", {})
        if criteria:
            crit_table = doc.add_table(rows=1, cols=3)
            crit_table.style = "Light List Accent 1"
            hdr = crit_table.rows[0].cells
            hdr[0].text = "Criterion"
            hdr[1].text = "Score"
            hdr[2].text = "Notes"
            for cell in hdr:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True

            for crit_id, crit_result in criteria.items():
                if isinstance(crit_result, dict):
                    row = crit_table.add_row()
                    # Use the criterion name or ID
                    crit_name = crit_id.replace("_", " ").replace("profit ", "").replace("empower ", "").replace("risk ", "").replace("board ", "").replace("value ", "").title()
                    row.cells[0].text = crit_name
                    crit_score = crit_result.get("score", 0)
                    row.cells[1].text = f"{crit_score}/2"
                    explanation = crit_result.get("explanation", "")
                    if len(explanation) > 200:
                        explanation = explanation[:197] + "..."
                    row.cells[2].text = explanation

        doc.add_paragraph()

    # Statute concerns
    statute_concerns = analysis_results.get("statute_concerns", [])
    if statute_concerns:
        doc.add_heading("Statutory Compliance Notes", level=1)
        for concern in statute_concerns:
            doc.add_paragraph(concern, style="List Bullet")
        doc.add_paragraph()

    # Recommendation
    doc.add_heading("Recommendation", level=1)
    recommendation = analysis_results.get("overall_recommendation", "")
    if recommendation:
        doc.add_paragraph(recommendation)
    else:
        doc.add_paragraph(
            "Based on our evaluation of your current agreement, we recommend "
            "a revised or full contract replacement to better align with best practices "
            "in profitability, empowerment, risk mitigation, and long-term value creation."
        )

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("This scorecard is provided for informational purposes and does not constitute legal advice.")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.save(output_path)
    return output_path


def _score_color(score):
    """Return RGB color based on score."""
    if score >= 0.7:
        return RGBColor(0, 128, 0)  # Green
    elif score >= 0.4:
        return RGBColor(204, 153, 0)  # Amber
    else:
        return RGBColor(192, 0, 0)  # Red


def _score_rating(score):
    """Return text rating based on score."""
    if score >= 0.8:
        return "Strong"
    elif score >= 0.6:
        return "Adequate"
    elif score >= 0.4:
        return "Needs Improvement"
    else:
        return "Weak"


def _color_heading(heading, color):
    """Apply color to all runs in a heading."""
    for run in heading.runs:
        run.font.color.rgb = color
